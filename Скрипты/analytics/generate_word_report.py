#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate professional Word report comparing DU processing scenarios.

Creates a comprehensive analysis document with:
- Executive summary
- Scenario comparison charts
- Statistical analysis
- Acceleration metrics
- Conclusions and recommendations
"""

import io
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

# Set style
sns.set_style("whitegrid")
plt.rcParams['font.family'] = 'Arial'
plt.rcParams['font.size'] = 10
plt.rcParams['figure.dpi'] = 150

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_PATH = PROJECT_ROOT / "data" / "du_tasks_times.csv"
OUT_PATH = PROJECT_ROOT / "Документация" / "Reports" / "du_speed_analysis" / "DU_Performance_Analysis.docx"


def read_data():
    """Read and prepare data."""
    df = pd.read_csv(DATA_PATH, sep=';', encoding='utf-8')
    df['date'] = pd.to_datetime(df['date'])
    return df


def create_chart_daily_comparison(df):
    """Chart 1: Daily comparison by scenario."""
    fig, ax = plt.subplots(figsize=(12, 6))
    
    colors = {'Типовой': '#FF8C42', 
              'Без дублей обменов': '#9B59B6',
              'Без дублей обменов + Параллельные портфели': '#16A085'}
    
    for scenario in df['scenario'].unique():
        data = df[df['scenario'] == scenario].sort_values('date')
        ax.bar(data['date'], data['minutes'], label=scenario, 
               color=colors.get(scenario, '#3498DB'), alpha=0.7, width=0.8)
    
    ax.set_xlabel('Дата', fontsize=12, fontweight='bold')
    ax.set_ylabel('Длительность (минуты)', fontsize=12, fontweight='bold')
    ax.set_title('Ежедневное время обработки по сценариям', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.legend(loc='upper right', fontsize=10)
    ax.grid(True, alpha=0.3)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf


def create_chart_box_comparison(df):
    """Chart 2: Box plot comparison."""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    scenarios_order = ['Типовой', 'Без дублей обменов', 
                       'Без дублей обменов + Параллельные портфели']
    colors_palette = ['#FF8C42', '#9B59B6', '#16A085']
    
    data_for_box = [df[df['scenario'] == s]['minutes'].values 
                    for s in scenarios_order if s in df['scenario'].unique()]
    labels = [s for s in scenarios_order if s in df['scenario'].unique()]
    
    bp = ax.boxplot(data_for_box, labels=labels, patch_artist=True, 
                    widths=0.6, showmeans=True)
    
    for patch, color in zip(bp['boxes'], colors_palette):
        patch.set_facecolor(color)
        patch.set_alpha(0.7)
    
    ax.set_ylabel('Длительность (минуты)', fontsize=12, fontweight='bold')
    ax.set_title('Сравнение распределения длительности по сценариям', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.grid(True, alpha=0.3, axis='y')
    plt.xticks(rotation=15, ha='right')
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf


def create_chart_rolling_average(df):
    """Chart 3: Rolling average by scenario."""
    fig, ax = plt.subplots(figsize=(12, 6))
    
    colors = {'Типовой': '#FF8C42', 
              'Без дублей обменов': '#9B59B6',
              'Без дублей обменов + Параллельные портфели': '#16A085'}
    
    for scenario in df['scenario'].unique():
        data = df[df['scenario'] == scenario].sort_values('date')
        roll = data.set_index('date')['minutes'].rolling(7, min_periods=3).mean()
        ax.plot(roll.index, roll.values, label=scenario, 
                color=colors.get(scenario, '#3498DB'), linewidth=2.5, marker='o', markersize=4)
    
    ax.set_xlabel('Дата', fontsize=12, fontweight='bold')
    ax.set_ylabel('Среднее за 7 дней (минуты)', fontsize=12, fontweight='bold')
    ax.set_title('Скользящее среднее (7 дней) по сценариям', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.legend(loc='upper right', fontsize=10)
    ax.grid(True, alpha=0.3)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf


def create_chart_heatmap_by_week(df):
    """Chart 4: Heatmap by week and scenario."""
    df_copy = df.copy()
    df_copy['week'] = pd.to_datetime(df_copy['date']).dt.isocalendar().week
    df_copy['year'] = pd.to_datetime(df_copy['date']).dt.year
    df_copy['week_label'] = df_copy['year'].astype(str) + '-W' + df_copy['week'].astype(str).str.zfill(2)
    
    pivot = df_copy.pivot_table(index='week_label', columns='scenario', values='minutes', aggfunc='mean')
    
    fig, ax = plt.subplots(figsize=(12, 8))
    sns.heatmap(pivot, annot=True, fmt='.0f', cmap='RdYlGn_r', cbar_kws={'label': 'Минуты'}, 
                linewidths=0.5, ax=ax)
    ax.set_title('Тепловая карта: среднее время по неделям и сценариям', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.set_xlabel('Сценарий', fontsize=12, fontweight='bold')
    ax.set_ylabel('Неделя', fontsize=12, fontweight='bold')
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf


def create_chart_summary_bars(df):
    """Chart 5: Summary comparison bar chart."""
    summary = df.groupby('scenario')['minutes'].agg(['mean', 'median', 'std']).reset_index()
    
    fig, ax = plt.subplots(figsize=(10, 6))
    
    scenarios = summary['scenario'].values
    x = np.arange(len(scenarios))
    width = 0.25
    
    colors = ['#3498DB', '#E74C3C', '#2ECC71']
    
    ax.bar(x - width, summary['mean'], width, label='Среднее', color=colors[0], alpha=0.8)
    ax.bar(x, summary['median'], width, label='Медиана', color=colors[1], alpha=0.8)
    ax.bar(x + width, summary['std'], width, label='Ст. отклонение', color=colors[2], alpha=0.8)
    
    ax.set_ylabel('Минуты', fontsize=12, fontweight='bold')
    ax.set_title('Статистическое сравнение сценариев', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(scenarios, rotation=15, ha='right')
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3, axis='y')
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf


def calculate_statistics(df):
    """Calculate comprehensive statistics."""
    stats = {}
    
    for scenario in df['scenario'].unique():
        data = df[df['scenario'] == scenario]['minutes']
        stats[scenario] = {
            'count': len(data),
            'mean': data.mean(),
            'median': data.median(),
            'std': data.std(),
            'min': data.min(),
            'max': data.max(),
            'total': data.sum()
        }
    
    # Calculate acceleration
    base = 'Типовой'
    if base in stats:
        base_mean = stats[base]['mean']
        for scenario in stats:
            if scenario != base:
                stats[scenario]['acceleration'] = base_mean / stats[scenario]['mean']
    
    return stats


def add_heading(doc, text, level=1):
    """Add styled heading."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return heading


def add_styled_paragraph(doc, text, bold=False, color=None):
    """Add styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def create_word_document(df, stats):
    """Create comprehensive Word document."""
    doc = Document()
    
    # Title page
    title = doc.add_heading('Анализ производительности обработки объектов после загрузки из DU', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Сравнительный анализ трёх сценариев')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(70, 70, 70)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f'Дата отчёта: {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(11)
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Executive Summary
    add_heading(doc, '1. Резюме', level=1)
    add_styled_paragraph(doc, 
        'Настоящий отчёт представляет комплексный анализ производительности системы обработки '
        'объектов после загрузки из DU за период июль-октябрь 2025 года. Проанализированы три сценария '
        'с прогрессивными улучшениями.')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Ключевые результаты:', bold=True)
    
    base_scenario = 'Типовой'
    if base_scenario in stats:
        base_mean = stats[base_scenario]['mean']
        for scenario in stats:
            if 'acceleration' in stats[scenario]:
                improvement = (1 - 1/stats[scenario]['acceleration']) * 100
                add_styled_paragraph(doc, 
                    f'• {scenario}: среднее время {stats[scenario]["mean"]:.1f} мин/день, '
                    f'ускорение в {stats[scenario]["acceleration"]:.1f}× раз '
                    f'(улучшение на {improvement:.0f}%)')
    
    doc.add_page_break()
    
    # Scenarios description
    add_heading(doc, '2. Описание сценариев', level=1)
    
    add_heading(doc, '2.1. Типовой (базовый)', level=2)
    add_styled_paragraph(doc, 
        f'Период: 01.07.2025 - 08.10.2025 ({stats[base_scenario]["count"]} дней)')
    add_styled_paragraph(doc, 
        'Стандартный сценарий обработки без оптимизаций. Используется как базовая линия для сравнения.')
    
    add_heading(doc, '2.2. Без дублей обменов', level=2)
    scenario2 = 'Без дублей обменов'
    if scenario2 in stats:
        add_styled_paragraph(doc, 
            f'Период: 09.10.2025 - 16.10.2025 ({stats[scenario2]["count"]} дней)')
        add_styled_paragraph(doc, 
            'Оптимизирован процесс обмена данными с устранением дублирующих операций.')
    
    add_heading(doc, '2.3. Без дублей обменов + Параллельные портфели', level=2)
    scenario3 = 'Без дублей обменов + Параллельные портфели'
    if scenario3 in stats:
        add_styled_paragraph(doc, 
            f'Период: 17.10.2025 - 30.10.2025 ({stats[scenario3]["count"]} дней)')
        add_styled_paragraph(doc, 
            'Добавлена параллельная обработка портфелей в дополнение к устранению дублей.')
    
    doc.add_page_break()
    
    # Charts section
    add_heading(doc, '3. Визуальный анализ', level=1)
    
    add_heading(doc, '3.1. Ежедневное время обработки', level=2)
    chart1 = create_chart_daily_comparison(df)
    doc.add_picture(chart1, width=Inches(6.5))
    add_styled_paragraph(doc, 
        'График показывает ежедневное время обработки для каждого сценария. '
        'Отчётливо видно снижение времени обработки при переходе между сценариями.')
    
    doc.add_page_break()
    
    add_heading(doc, '3.2. Статистическое распределение', level=2)
    chart2 = create_chart_box_comparison(df)
    doc.add_picture(chart2, width=Inches(6.5))
    add_styled_paragraph(doc, 
        'Box plot демонстрирует распределение значений, медиану и выбросы для каждого сценария.')
    
    doc.add_page_break()
    
    add_heading(doc, '3.3. Скользящее среднее', level=2)
    chart3 = create_chart_rolling_average(df)
    doc.add_picture(chart3, width=Inches(6.5))
    add_styled_paragraph(doc, 
        '7-дневное скользящее среднее показывает общие тренды без краткосрочных колебаний.')
    
    doc.add_page_break()
    
    add_heading(doc, '3.4. Коэффициент ускорения', level=2)
    chart4 = create_chart_acceleration(df)
    doc.add_picture(chart4, width=Inches(6.5))
    add_styled_paragraph(doc, 
        'График демонстрирует во сколько раз каждый сценарий быстрее базового.')
    
    doc.add_page_break()
    
    add_heading(doc, '3.5. Сводное сравнение', level=2)
    chart5 = create_chart_summary_bars(df)
    doc.add_picture(chart5, width=Inches(6.5))
    add_styled_paragraph(doc, 
        'Сравнение ключевых статистических показателей: среднее, медиана и стандартное отклонение.')
    
    doc.add_page_break()
    
    # Statistical table
    add_heading(doc, '4. Статистические показатели', level=1)
    
    table = doc.add_table(rows=len(stats)+1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Header
    headers = ['Сценарий', 'Дней', 'Среднее', 'Медиана', 'Мин', 'Макс', 'Ст.откл.', 'Ускорение']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for idx, (scenario, stat) in enumerate(stats.items(), 1):
        row = table.rows[idx]
        row.cells[0].text = scenario
        row.cells[1].text = str(stat['count'])
        row.cells[2].text = f"{stat['mean']:.1f}"
        row.cells[3].text = f"{stat['median']:.1f}"
        row.cells[4].text = f"{stat['min']:.0f}"
        row.cells[5].text = f"{stat['max']:.0f}"
        row.cells[6].text = f"{stat['std']:.1f}"
        row.cells[7].text = f"{stat.get('acceleration', 1.0):.2f}×"
    
    doc.add_page_break()
    
    # Conclusions
    add_heading(doc, '5. Выводы и рекомендации', level=1)
    
    add_styled_paragraph(doc, '5.1. Основные выводы:', bold=True)
    add_styled_paragraph(doc, 
        '• Внедрение оптимизаций привело к существенному улучшению производительности системы')
    add_styled_paragraph(doc, 
        '• Финальный сценарий демонстрирует стабильную работу с минимальным временем обработки')
    add_styled_paragraph(doc, 
        '• Наибольший эффект достигнут за счёт комбинации устранения дублей и распараллеливания')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, '5.2. Рекомендации:', bold=True)
    add_styled_paragraph(doc, 
        '• Продолжить мониторинг производительности в текущей конфигурации')
    add_styled_paragraph(doc, 
        '• Рассмотреть возможность дополнительной оптимизации для дней с пиковой нагрузкой')
    add_styled_paragraph(doc, 
        '• Документировать лучшие практики для использования в других подсистемах')
    
    return doc


def main():
    """Main execution."""
    try:
        print("📊 Загрузка данных...")
        df = read_data()
        
        print("📈 Расчёт статистики...")
        stats = calculate_statistics(df)
        
        print("📝 Создание Word документа...")
        doc = create_word_document(df, stats)
        
        OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_PATH)
        
        print(f"\n✅ Отчёт успешно создан: {OUT_PATH}")
        print(f"\nСтатистика по сценариям:")
        for scenario, stat in stats.items():
            print(f"\n{scenario}:")
            print(f"  Дней: {stat['count']}")
            print(f"  Среднее: {stat['mean']:.1f} мин")
            print(f"  Медиана: {stat['median']:.1f} мин")
            if 'acceleration' in stat:
                print(f"  Ускорение: {stat['acceleration']:.2f}×")
        
        return 0
        
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    raise SystemExit(main())

